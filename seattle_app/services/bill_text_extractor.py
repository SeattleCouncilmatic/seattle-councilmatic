"""Download and extract plain text from bill attachments.

Seattle bills carry their substantive content in Legistar document
attachments (PDF or .docx). The :mod:`Bill` row from the OCD/pupa scrape
only stores their URLs, not their content; this module downloads them
on demand and turns them into plain text suitable for LLM
summarization.

Two well-known attachment kinds — see the ``extract_bill_text``
management command for the picker that decides which to use:

* "Summary and Fiscal Note" (.docx) — staff plain-language summary plus
  fiscal analysis. Available from introduction onward.
* "Signed Ordinance NNNNN" / "Signed Resolution NNNNN" (PDF) — canonical
  body text. Only present after enactment.

We skip the "Affidavit of Publication" (PDF, legal notice, no bill
content). Other attachment types fall to ``other`` and are ignored by
default; callers can opt in with ``include_other=True`` if they want a
broader sweep.
"""
from __future__ import annotations

import io
import logging
import re
import signal
from contextlib import contextmanager
from dataclasses import dataclass
from typing import Iterable, Optional

import pdfplumber
import requests
from docx import Document

logger = logging.getLogger(__name__)

# Conservative HTTP timeout — Legistar's CDN is fast under normal load,
# but a stuck request shouldn't block a whole batch run.
_HTTP_TIMEOUT_SECONDS = 30

# Bytes ceiling for a single attachment download. Most "Full Text" /
# "Summary and Fiscal Note" docs are well under 1 MB; some bills attach
# multi-MB EIS reports or comprehensive-plan PDFs that pdfplumber will
# happily try to parse and balloon memory on. Cap at 50 MB and skip
# anything larger.
_MAX_DOWNLOAD_BYTES = 50 * 1024 * 1024  # 50 MB

# How long an extracted blob can grow before we treat it as suspicious
# and bail out. Bill texts in the wild run 5k–80k chars; anything past
# this is almost certainly a scan/OCR artifact or wrong document.
_MAX_TEXT_CHARS = 500_000

# Read chunk size when streaming the download. Smaller chunks help us
# notice we've exceeded _MAX_DOWNLOAD_BYTES quickly without buffering
# more than necessary.
_DOWNLOAD_CHUNK_BYTES = 64 * 1024

# Page-count cap on PDF extraction. Most bill attachments are under
# 50 pages; "Signed Ordinance" PDFs that embed full settlement
# agreements or comprehensive-plan documents can run into the
# hundreds. pdfplumber's per-page parse cost adds up faster than the
# per-page cache flush releases memory, so we hard-skip past this
# threshold rather than spending minutes (and gigabytes of RSS) on
# what's almost never the canonical bill body anyway.
_MAX_PDF_PAGES = 500

# Wall-clock cap on a single PDF parse. SIGALRM-based, so it works
# inside the Linux Docker container (manage.py runs in the main
# thread). Catches malformed PDFs that send pdfplumber/pdfminer into
# an effectively-infinite loop, and over-large legitimate PDFs whose
# page count slipped under the cap but still take forever.
_PDF_EXTRACT_TIMEOUT_SECONDS = 90

# Document categories. Matched against BillDocument.note (case-insensitive,
# anchored). Order here mirrors the LLM-input concatenation order.
SUMMARY_NOTE_RE = re.compile(r"^summary\s+and\s+fiscal\s+note\b", re.IGNORECASE)
# Permissive prefix match: "Signed Ordinance 127119", "Signed Resolution
# 32168", "Signed Council Bill 12345" all qualify. Risk of catching a
# stray "Signed [Something Else]" doc is small in practice — Legistar's
# document templates are predictable.
SIGNED_NOTE_RE = re.compile(r"^signed\s+", re.IGNORECASE)
# Pre-enactment canonical text. Legistar attaches the ordinance body as
# "Full Text: CB 121173 v1" (or vN for revised drafts) BEFORE the bill
# is signed; once signed, it gets renamed and re-attached as "Signed
# Ordinance NNNNN". Both are canonical bill body for our summarization
# purposes — categorizing separately just so the audit trail can tell
# them apart and so future code can prefer signed over draft if both
# are present.
FULL_TEXT_NOTE_RE = re.compile(r"^full\s+text\b", re.IGNORECASE)
AFFIDAVIT_NOTE_RE = re.compile(r"\baffidavit\b", re.IGNORECASE)


# Media-type strings, kept verbose for legibility.
_PDF_MEDIA = "application/pdf"
_DOCX_MEDIA = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_DOC_MEDIA = "application/msword"  # legacy binary; we don't extract from these


@dataclass
class ExtractedDocument:
    """Result of extracting one document, including everything we want
    to record on the BillText row's ``source_documents`` JSON."""

    note: str
    url: str
    media_type: str
    category: str          # 'summary' | 'signed' | 'affidavit' | 'other' | 'unsupported'
    text: str              # may be empty for skipped/failed docs
    error: Optional[str] = None


def categorize_note(note: str) -> str:
    """Classify a BillDocument note into our document buckets."""
    if SUMMARY_NOTE_RE.match(note or ""):
        return "summary"
    if SIGNED_NOTE_RE.match(note or ""):
        return "signed"
    if FULL_TEXT_NOTE_RE.match(note or ""):
        return "full_text"
    if AFFIDAVIT_NOTE_RE.search(note or ""):
        return "affidavit"
    return "other"


def extract_text(url: str, media_type: str) -> str:
    """Download `url` and return its plain text.

    Returns "" on download failure, unsupported media type (legacy .doc),
    download exceeding _MAX_DOWNLOAD_BYTES, or extraction failure.
    Errors are logged at WARNING; callers decide whether to skip the
    document or surface the error.

    The download is streamed and capped to bound memory: we abort if
    Content-Length advertises more than the cap, and we stop reading
    once we've accumulated more than the cap regardless of headers
    (some servers don't send Content-Length, or lie about it).
    """
    # Skip the whole download for unsupported media types — no point
    # eating bytes on a legacy .doc we'll just throw out.
    if media_type == _DOC_MEDIA:
        logger.warning("skipping legacy .doc (not supported): %s", url)
        return ""
    if media_type not in (_PDF_MEDIA, _DOCX_MEDIA):
        logger.warning("unsupported media type %r for %s", media_type, url)
        return ""

    try:
        with requests.get(url, timeout=_HTTP_TIMEOUT_SECONDS, stream=True) as resp:
            resp.raise_for_status()
            advertised = resp.headers.get("Content-Length")
            if advertised and advertised.isdigit() and int(advertised) > _MAX_DOWNLOAD_BYTES:
                logger.warning(
                    "skipping oversized download (Content-Length=%s, cap=%d): %s",
                    advertised, _MAX_DOWNLOAD_BYTES, url,
                )
                return ""
            buf = bytearray()
            for chunk in resp.iter_content(chunk_size=_DOWNLOAD_CHUNK_BYTES):
                if not chunk:
                    continue
                buf.extend(chunk)
                if len(buf) > _MAX_DOWNLOAD_BYTES:
                    logger.warning(
                        "aborting download — exceeded %d bytes: %s",
                        _MAX_DOWNLOAD_BYTES, url,
                    )
                    return ""
            blob = bytes(buf)
    except requests.RequestException as e:
        logger.warning("download failed: %s — %s", url, e)
        return ""

    if media_type == _PDF_MEDIA:
        return _extract_pdf(blob, url)
    if media_type == _DOCX_MEDIA:
        return _extract_docx(blob, url)
    return ""


def combine_bill_documents(
    documents: Iterable[dict],
    *,
    include_other: bool = False,
    progress=None,
) -> tuple[str, list[ExtractedDocument]]:
    """Pick the right documents for a bill, extract each, and return a
    single concatenated text block plus a per-document audit trail.

    Each document in `documents` is a dict with ``note``, ``url``, and
    ``media_type`` keys (matching the shape we already serialize in the
    bill detail API).

    `progress` is an optional callback ``(note, category, status, chars)``
    invoked at each per-document state transition. ``status`` is one
    of ``"extracting"`` (about to download/parse), ``"done"`` (text
    extracted; chars is the count), ``"skipped"`` (category was
    affidavit/other or text was empty/too large). Use it to stream
    progress in CLI tools without refactoring this helper into a
    generator.

    The returned text is structured for LLM consumption:

        [STAFF SUMMARY AND FISCAL NOTE — note text]
        <body>

        [SIGNED ORDINANCE NNNNN — note text]
        <body>

    Section markers help the LLM tell staff framing apart from the
    canonical text. If a category yields no text, its block is omitted.
    """
    def _emit(note: str, category: str, status: str, chars: int = 0) -> None:
        if progress is not None:
            progress(note, category, status, chars)

    extracted: list[ExtractedDocument] = []
    for doc in documents:
        note = (doc.get("note") or "").strip()
        url = (doc.get("url") or "").strip()
        media_type = (doc.get("media_type") or "").strip()
        category = categorize_note(note)

        if category == "affidavit":
            # Legal notice; never has bill content. Record it for the
            # audit trail but don't download.
            _emit(note, category, "skipped")
            extracted.append(ExtractedDocument(
                note=note, url=url, media_type=media_type,
                category=category, text="",
            ))
            continue
        if category == "other" and not include_other:
            _emit(note, category, "skipped")
            extracted.append(ExtractedDocument(
                note=note, url=url, media_type=media_type,
                category=category, text="",
            ))
            continue
        if not url:
            _emit(note, category, "skipped")
            extracted.append(ExtractedDocument(
                note=note, url=url, media_type=media_type,
                category=category, text="",
                error="missing url",
            ))
            continue

        _emit(note, category, "extracting")
        text = extract_text(url, media_type)
        if len(text) > _MAX_TEXT_CHARS:
            err = f"extracted text too large ({len(text)} chars), suspicious"
            logger.warning("%s: %s", url, err)
            _emit(note, category, "skipped")
            extracted.append(ExtractedDocument(
                note=note, url=url, media_type=media_type,
                category=category, text="", error=err,
            ))
            continue
        _emit(note, category, "done", len(text))
        extracted.append(ExtractedDocument(
            note=note, url=url, media_type=media_type,
            category=category, text=text,
        ))

    parts: list[str] = []
    # Summary first (staff framing), then canonical text (signed if
    # available, else the pre-enactment "Full Text" draft). Multiple
    # matches in a category get all included; for an enacted bill that
    # also still has a draft attached, the LLM sees both — small cost
    # in tokens, robust to versions diverging.
    sections = (
        ("summary",   "STAFF SUMMARY AND FISCAL NOTE"),
        ("signed",    "SIGNED CANONICAL TEXT"),
        ("full_text", "FULL TEXT (PRE-ENACTMENT DRAFT)"),
    )
    for category, header_word in sections:
        for doc in extracted:
            if doc.category == category and doc.text:
                parts.append(f"[{header_word} — {doc.note}]\n{doc.text}")
    return "\n\n".join(parts), extracted


# ---------------------------------------------------------------------------
# Format-specific extraction
# ---------------------------------------------------------------------------

class _ExtractionTimeout(Exception):
    """Internal — raised by SIGALRM handler when a PDF parse runs too long."""


@contextmanager
def _pdf_extract_timeout(seconds: int, url: str):
    """Set a SIGALRM-based wall-clock timeout for the enclosed block.

    Linux-only (signal.SIGALRM doesn't exist on Windows). The Docker
    container runs Linux and manage.py runs in the main thread, so this
    works for our use case. Falls back to a no-op on non-Linux callers
    (e.g. running tests on Windows directly).

    Note: signals only interrupt at Python bytecode boundaries, so
    pdfplumber blocking deep in a C extension might not return
    immediately. In practice pdfplumber/pdfminer is mostly Python and
    interrupts cleanly within a few hundred ms.
    """
    if not hasattr(signal, "SIGALRM"):
        yield
        return

    def _handler(signum, frame):
        raise _ExtractionTimeout(f"pdf extraction exceeded {seconds}s: {url}")

    prev = signal.signal(signal.SIGALRM, _handler)
    signal.alarm(seconds)
    try:
        yield
    finally:
        signal.alarm(0)
        signal.signal(signal.SIGALRM, prev)


def _extract_pdf(blob: bytes, url: str) -> str:
    """Plain-text extraction from a PDF blob via pdfplumber.

    pdfplumber's ``extract_text()`` is single-column-friendly and
    suitable for the well-formatted ordinance PDFs Legistar produces.
    Page-by-page join with single newlines so paragraph reflow is
    tractable downstream.

    Two memory/runtime bounds:

    1. **Page count cap.** PDFs with more than ``_MAX_PDF_PAGES`` pages
       are skipped without per-page extraction. Bill bodies don't run
       300+ pages; an attached comprehensive-plan EIS or settlement
       agreement does, and that's not the canonical bill text we want.
    2. **Wall-clock timeout.** Each parse is bounded by
       ``_PDF_EXTRACT_TIMEOUT_SECONDS`` via SIGALRM. Catches malformed
       PDFs that send pdfminer into pathological loops, and otherwise
       gracefully aborts on any single doc taking unreasonable time.

    Each page's parsed-char cache is also flushed immediately after
    extraction so memory holds at most one page's worth of intermediate
    state at any given time.
    """
    try:
        with _pdf_extract_timeout(_PDF_EXTRACT_TIMEOUT_SECONDS, url):
            with pdfplumber.open(io.BytesIO(blob)) as pdf:
                num_pages = len(pdf.pages)
                if num_pages > _MAX_PDF_PAGES:
                    logger.warning(
                        "skipping PDF with too many pages (%d > %d): %s",
                        num_pages, _MAX_PDF_PAGES, url,
                    )
                    return ""
                pages: list[str] = []
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(text)
                    page.flush_cache()
            return "\n".join(pages).strip()
    except _ExtractionTimeout as e:
        logger.warning("%s", e)
        return ""
    except Exception as e:
        logger.warning("pdfplumber failed on %s: %s", url, e)
        return ""


def _extract_docx(blob: bytes, url: str) -> str:
    """Plain-text extraction from a .docx blob via python-docx.

    Walks paragraphs in document order and tables row-by-row.
    Headers/footers are ignored — Seattle's "Summary and Fiscal Note"
    template puts the actual content in the body, with template
    cruft in headers/footers we don't want polluting the LLM input.
    """
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as e:
        logger.warning("python-docx failed on %s: %s", url, e)
        return ""

    parts: list[str] = []
    # Walk the document body. Element order matters because a "Summary
    # and Fiscal Note" template intersperses headings, paragraphs, and
    # tables; reading them in document order preserves intent.
    seen_table_ids: set[int] = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        # Avoid double-counting tables that appear inside body
        # paragraphs; python-docx returns top-level tables here.
        if id(table) in seen_table_ids:
            continue
        seen_table_ids.add(id(table))
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
